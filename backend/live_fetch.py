from __future__ import annotations

import os
import re
import threading
import time
import logging
from io import BytesIO
from pathlib import PurePosixPath
from urllib.parse import unquote, urljoin, urlparse

import requests
from bs4 import BeautifulSoup
from ocr_client import (
    is_paddle_ocr_enabled,
    is_supported_ocr_image_extension,
    run_paddle_ocr_on_image,
    run_paddle_ocr_on_pdf,
    should_run_pdf_ocr,
)
from requests.adapters import HTTPAdapter

try:
    from docx import Document
except ModuleNotFoundError:
    Document = None

try:
    from pypdf import PdfReader
except ModuleNotFoundError:
    PdfReader = None

logging.getLogger("pypdf").setLevel(logging.ERROR)
HEADERS = {"User-Agent": "UVT_Asist/1.0"}


def env_int(name: str, default: str, minimum: int) -> int:
    try:
        value = int(os.getenv(name, default))
    except (TypeError, ValueError):
        value = int(default)
    return max(minimum, value)


MAX_TEXT_LENGTH = env_int("LIVE_FETCH_MAX_TEXT_LENGTH", "16000", 4000)
MAX_PDF_PAGES = env_int("LIVE_FETCH_MAX_PDF_PAGES", "12", 1)
INDEX_HTML_MAX_TEXT_LENGTH = max(
    MAX_TEXT_LENGTH,
    env_int("INDEX_HTML_FETCH_TEXT_CHARS", "24000", MAX_TEXT_LENGTH),
)
INDEX_DOCUMENT_MAX_TEXT_LENGTH = max(
    MAX_TEXT_LENGTH,
    env_int("INDEX_DOCUMENT_FETCH_TEXT_CHARS", "120000", MAX_TEXT_LENGTH),
)
INDEX_DOCUMENT_MAX_PDF_PAGES = max(
    MAX_PDF_PAGES,
    env_int("INDEX_DOCUMENT_FETCH_PDF_PAGES", "80", MAX_PDF_PAGES),
)
CACHE_TTL = max(60, int(os.getenv("LIVE_FETCH_CACHE_TTL", "300")))

BAD_EXTENSIONS = (".gif", ".svg", ".zip", ".rar", ".doc", ".xls", ".xlsx", ".ppt", ".pptx")
NON_HTML_LINK_PREFIXES = ("#", "mailto:", "tel:", "javascript:")

REMOVABLE_SELECTORS = (
    "header",
    "nav",
    "aside",
    "form",
    "iframe",
    "button",
    "input",
    "select",
    "textarea",
    "[role='navigation']",
    ".menu",
    ".nav",
    ".navbar",
    ".breadcrumbs",
    ".breadcrumb",
    ".cookie",
    ".cookies",
    ".gdpr",
    ".banner",
    ".sidebar",
    ".search-form",
    ".newsletter",
    "#cookie-notice",
    "#cookie-law-info-bar",
)

PAGE_CACHE: dict[str, dict] = {}
LINK_CACHE: dict[str, dict] = {}
CACHE_LOCK = threading.Lock()
SESSION_LOCAL = threading.local()


def get_session() -> requests.Session:
    session = getattr(SESSION_LOCAL, "session", None)
    if session is not None:
        return session

    session = requests.Session()
    adapter = HTTPAdapter(pool_connections=16, pool_maxsize=16, max_retries=0)
    session.mount("http://", adapter)
    session.mount("https://", adapter)
    session.headers.update(HEADERS)
    SESSION_LOCAL.session = session
    return session


def get_cached(cache: dict, key: str):
    now = time.time()
    with CACHE_LOCK:
        cached = cache.get(key)
        if cached and now - cached["timestamp"] < CACHE_TTL:
            return cached["data"]
    return None


def set_cached(cache: dict, key: str, data) -> None:
    with CACHE_LOCK:
        cache[key] = {"timestamp": time.time(), "data": data}


def clear_fetch_caches() -> None:
    with CACHE_LOCK:
        PAGE_CACHE.clear()
        LINK_CACHE.clear()


def normalize_whitespace(text: str) -> str:
    return re.sub(r"\s+", " ", str(text)).strip()


def normalize_hostname(hostname: str) -> str:
    hostname = hostname.strip().lower()
    return hostname[4:] if hostname.startswith("www.") else hostname


def same_domain(url: str, base_urls: list[str]) -> bool:
    host = normalize_hostname(urlparse(url).hostname or "")
    return any(host == normalize_hostname(urlparse(base_url).hostname or "") for base_url in base_urls)


def get_url_extension(url: str) -> str:
    return PurePosixPath(urlparse(url).path.lower()).suffix


def truncate_text(text: str, max_length: int = MAX_TEXT_LENGTH) -> str:
    return normalize_whitespace(text)[:max(0, int(max_length or MAX_TEXT_LENGTH))]


def is_html_response(content_type: str, url: str) -> bool:
    content_type = (content_type or "").lower()
    return "html" in content_type or get_url_extension(url) in {"", ".html", ".htm", ".php", ".aspx"}


def is_pdf_response(content_type: str, url: str) -> bool:
    content_type = (content_type or "").lower()
    return "application/pdf" in content_type or get_url_extension(url) == ".pdf"


def looks_like_pdf_content(content: bytes) -> bool:
    return bytes(content or b"").lstrip().startswith(b"%PDF-")


def is_docx_response(content_type: str, url: str) -> bool:
    content_type = (content_type or "").lower()
    return (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document" in content_type
        or get_url_extension(url) == ".docx"
    )


def is_text_response(content_type: str, url: str) -> bool:
    content_type = (content_type or "").lower()
    return content_type.startswith("text/plain") or get_url_extension(url) == ".txt"


def is_ocr_image_response(content_type: str, url: str) -> bool:
    content_type = (content_type or "").lower()
    extension = get_url_extension(url)
    return content_type.startswith(("image/jpeg", "image/png")) or is_supported_ocr_image_extension(extension)


def clean_html_text(html: str) -> str:
    soup = BeautifulSoup(html, "html.parser")

    for tag in soup(["script", "style", "noscript", "svg", "img", "footer", "template"]):
        tag.decompose()

    for selector in REMOVABLE_SELECTORS:
        for node in soup.select(selector):
            node.decompose()

    return normalize_whitespace(soup.get_text(separator=" "))


def extract_pdf_text(
    content: bytes,
    max_text_length: int = MAX_TEXT_LENGTH,
    max_pdf_pages: int = MAX_PDF_PAGES,
) -> str:
    if PdfReader is None:
        return ""

    reader = PdfReader(BytesIO(content))
    pages: list[str] = []
    for page in reader.pages[:max(1, int(max_pdf_pages or MAX_PDF_PAGES))]:
        page_text = page.extract_text() or ""
        if page_text.strip():
            pages.append(page_text)
        if sum(len(item) for item in pages) >= max_text_length:
            break

    return truncate_text(" ".join(pages), max_text_length)


def extract_pdf_text_with_fallback(
    content: bytes,
    max_text_length: int = MAX_TEXT_LENGTH,
    max_pdf_pages: int = MAX_PDF_PAGES,
) -> str:
    text = extract_pdf_text(content, max_text_length=max_text_length, max_pdf_pages=max_pdf_pages)
    if should_run_pdf_ocr(text):
        ocr_text = run_paddle_ocr_on_pdf(content)
        if ocr_text:
            return truncate_text(ocr_text, max_text_length)
    return text


def extract_docx_text(content: bytes, max_text_length: int = MAX_TEXT_LENGTH) -> str:
    if Document is None:
        return ""

    document = Document(BytesIO(content))
    parts = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
    for table in document.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                parts.append(row_text)

    return truncate_text(" ".join(parts), max_text_length)


def extract_image_text(content: bytes, url: str, max_text_length: int = MAX_TEXT_LENGTH) -> str:
    if not is_paddle_ocr_enabled():
        return ""
    extension = get_url_extension(url) or ".png"
    return truncate_text(run_paddle_ocr_on_image(content, extension) or "", max_text_length)


def extract_filename_from_content_disposition(content_disposition: str) -> str:
    if not content_disposition:
        return ""

    encoded_match = re.search(r"filename\*\s*=\s*(?:UTF-8''|utf-8'')([^;]+)", content_disposition)
    if encoded_match:
        return unquote(encoded_match.group(1).strip().strip('"'))

    plain_match = re.search(r'filename\s*=\s*"?(?P<name>[^";]+)"?', content_disposition)
    return plain_match.group("name").strip() if plain_match else ""


def build_title_from_url(url: str, fallback: str = "") -> str:
    filename = extract_filename_from_content_disposition(fallback)
    if filename:
        return filename

    path_name = PurePosixPath(urlparse(url).path).name
    return path_name or fallback.strip() or url


def response_limits(
    content_type: str,
    final_url: str,
    index_mode: bool = False,
    max_text_length: int | None = None,
    max_pdf_pages: int | None = None,
) -> tuple[int, int]:
    document_response = (
        is_pdf_response(content_type, final_url)
        or is_docx_response(content_type, final_url)
        or is_text_response(content_type, final_url)
    )
    if index_mode and document_response:
        text_default = INDEX_DOCUMENT_MAX_TEXT_LENGTH
    elif index_mode:
        text_default = INDEX_HTML_MAX_TEXT_LENGTH
    else:
        text_default = MAX_TEXT_LENGTH
    pdf_default = INDEX_DOCUMENT_MAX_PDF_PAGES if index_mode and is_pdf_response(content_type, final_url) else MAX_PDF_PAGES
    return (
        max(1000, int(max_text_length or text_default)),
        max(1, int(max_pdf_pages or pdf_default)),
    )


def extract_response_text(
    response: requests.Response,
    final_url: str,
    max_text_length: int = MAX_TEXT_LENGTH,
    max_pdf_pages: int = MAX_PDF_PAGES,
) -> tuple[str, str, str]:
    content_type = response.headers.get("Content-Type", "")
    content_disposition = response.headers.get("Content-Disposition", "")

    title = build_title_from_url(final_url, fallback=content_disposition)
    if is_pdf_response(content_type, final_url) and not looks_like_pdf_content(response.content):
        return title, "", "invalid_pdf"

    if is_html_response(content_type, final_url):
        soup = BeautifulSoup(response.text, "html.parser")
        title = soup.title.get_text(strip=True) if soup.title else build_title_from_url(final_url)
        return title, truncate_text(clean_html_text(response.text), max_text_length), "html"

    if is_pdf_response(content_type, final_url):
        return title, extract_pdf_text_with_fallback(
            response.content,
            max_text_length=max_text_length,
            max_pdf_pages=max_pdf_pages,
        ), "pdf"
    if is_docx_response(content_type, final_url):
        return title, extract_docx_text(response.content, max_text_length=max_text_length), "docx"
    if is_text_response(content_type, final_url):
        return title, truncate_text(response.text, max_text_length), "text"
    if is_ocr_image_response(content_type, final_url):
        return title, extract_image_text(response.content, final_url, max_text_length=max_text_length), "image"

    return title, "", "unknown"


def fetch_page(
    url: str,
    timeout: int = 10,
    *,
    index_mode: bool = False,
    max_text_length: int | None = None,
    max_pdf_pages: int | None = None,
) -> dict:
    cache_key = (
        f"{url}|ocr={int(is_paddle_ocr_enabled())}|index={int(index_mode)}|"
        f"chars={max_text_length or ''}|pdf_pages={max_pdf_pages or ''}"
    )
    cached = get_cached(PAGE_CACHE, cache_key)
    if cached is not None:
        return cached

    try:
        response = get_session().get(url, timeout=(2, timeout))
        response.raise_for_status()
        final_url = response.url or url
        effective_text_length, effective_pdf_pages = response_limits(
            response.headers.get("Content-Type", ""),
            final_url,
            index_mode=index_mode,
            max_text_length=max_text_length,
            max_pdf_pages=max_pdf_pages,
        )
        title, text, response_type = extract_response_text(
            response,
            final_url,
            max_text_length=effective_text_length,
            max_pdf_pages=effective_pdf_pages,
        )
        page = {
            "url": final_url,
            "title": title,
            "text": truncate_text(text, effective_text_length),
            "type": response_type,
        }
    except Exception as exc:
        page = {
            "url": url,
            "title": url,
            "text": "",
            "error": str(exc),
            "type": "error",
        }

    set_cached(PAGE_CACHE, cache_key, page)
    return page


def _is_allowed_link(url: str, base_urls: list[str]) -> bool:
    parsed = urlparse(url)
    if parsed.scheme not in {"http", "https"}:
        return False
    if not same_domain(url, base_urls):
        return False

    extension = get_url_extension(url)
    if extension in BAD_EXTENSIONS:
        return False
    if extension in {".jpg", ".jpeg", ".png"} and not is_paddle_ocr_enabled():
        return False
    return True


def extract_candidate_links(start_url: str, base_urls: list[str], max_links: int = 20, timeout: int = 10) -> list[str]:
    cache_key = f"{start_url}|ocr={int(is_paddle_ocr_enabled())}"
    cached = get_cached(LINK_CACHE, cache_key)
    if cached is not None:
        return cached[:max_links]

    links: list[str] = []
    try:
        response = get_session().get(start_url, timeout=(2, timeout))
        response.raise_for_status()
        page_url = response.url or start_url
        if not is_html_response(response.headers.get("Content-Type", ""), page_url):
            return []

        soup = BeautifulSoup(response.text, "html.parser")
        for anchor in soup.select("a[href]"):
            href = (anchor.get("href") or "").strip()
            if not href or href.lower().startswith(NON_HTML_LINK_PREFIXES):
                continue

            full_url = urljoin(page_url, href).split("#")[0]
            if _is_allowed_link(full_url, base_urls) and full_url not in links:
                links.append(full_url)
            if len(links) >= max_links:
                break
    except Exception:
        links = []

    set_cached(LINK_CACHE, cache_key, links)
    return links[:max_links]
